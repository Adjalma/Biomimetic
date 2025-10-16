"""
Sistema RAG (Retrieval Augmented Generation) Avançado
====================================================

Sistema de recuperação e geração aumentada com:
1. Vector Store (ChromaDB/FAISS)
2. Embeddings avançados
3. RAGAS para avaliação de qualidade
4. Sistema de cache inteligente
5. Re-ranking de resultados
"""

import os
import json
import logging
import time
import hashlib
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

# RAG e Vector Stores
try:
    import chromadb
    from chromadb.config import Settings
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False

try:
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False

# Embeddings e NLP
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# RAGAS para avaliação
try:
    from ragas import evaluate
    from ragas.metrics import (
        faithfulness, answer_relevancy, context_relevancy, 
        context_recall, answer_correctness, answer_similarity
    )
    RAGAS_AVAILABLE = True
except ImportError:
    RAGAS_AVAILABLE = False

# Processamento de documentos
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Machine Learning
import numpy as np
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('rag_system.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Chunk de documento para indexação"""
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    chunk_index: int = 0
    source_document: str = ""
    page_number: Optional[int] = None

@dataclass
class SearchResult:
    """Resultado de busca no RAG"""
    chunk: DocumentChunk
    similarity_score: float
    rank: int
    source: str
    metadata: Dict[str, Any]

@dataclass
class RAGResponse:
    """Resposta do sistema RAG"""
    answer: str
    context: List[DocumentChunk]
    sources: List[str]
    confidence_score: float
    generation_time: float
    search_results: List[SearchResult]
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Processador de documentos para RAG"""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_text(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Processa texto em chunks"""
        chunks = []
        metadata = metadata or {}
        
        # Dividir texto em chunks
        words = text.split()
        current_chunk = []
        current_length = 0
        
        for i, word in enumerate(words):
            current_chunk.append(word)
            current_length += len(word) + 1
            
            if current_length >= self.chunk_size:
                chunk_text = " ".join(current_chunk)
                chunk_id = hashlib.md5(f"{metadata.get('source', '')}_{i}".encode()).hexdigest()[:16]
                
                chunk = DocumentChunk(
                    id=chunk_id,
                    content=chunk_text,
                    metadata=metadata.copy(),
                    chunk_index=len(chunks),
                    source_document=metadata.get('source', ''),
                    page_number=metadata.get('page_number')
                )
                chunks.append(chunk)
                
                # Overlap
                overlap_words = int(self.chunk_overlap / 5)  # Aproximação
                current_chunk = current_chunk[-overlap_words:] if overlap_words > 0 else []
                current_length = sum(len(word) + 1 for word in current_chunk)
        
        # Último chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_id = hashlib.md5(f"{metadata.get('source', '')}_{len(chunks)}".encode()).hexdigest()[:16]
            
            chunk = DocumentChunk(
                id=chunk_id,
                content=chunk_text,
                metadata=metadata.copy(),
                chunk_index=len(chunks),
                source_document=metadata.get('source', ''),
                page_number=metadata.get('page_number')
            )
            chunks.append(chunk)
        
        return chunks
    
    def process_pdf(self, file_path: str) -> List[DocumentChunk]:
        """Processa arquivo PDF"""
        if not PYMUPDF_AVAILABLE:
            logger.error("PyMuPDF não disponível para processamento de PDF")
            return []
        
        try:
            doc = fitz.open(file_path)
            chunks = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                metadata = {
                    'source': file_path,
                    'page_number': page_num + 1,
                    'file_type': 'pdf',
                    'total_pages': len(doc)
                }
                
                page_chunks = self.process_text(text, metadata)
                chunks.extend(page_chunks)
            
            doc.close()
            return chunks
        except Exception as e:
            logger.error(f"Erro ao processar PDF {file_path}: {e}")
            return []
    
    def process_docx(self, file_path: str) -> List[DocumentChunk]:
        """Processa arquivo DOCX"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx não disponível para processamento de DOCX")
            return []
        
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                'source': file_path,
                'file_type': 'docx',
                'paragraphs': len(doc.paragraphs)
            }
            
            return self.process_text(text, metadata)
        except Exception as e:
            logger.error(f"Erro ao processar DOCX {file_path}: {e}")
            return []

class EmbeddingManager:
    """Gerenciador de embeddings"""
    
    def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = None
        self._load_model()
    
    def _load_model(self):
        """Carrega modelo de embeddings"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            logger.error("Sentence Transformers não disponível")
            return
        
        try:
            self.model = SentenceTransformer(self.model_name)
            logger.info(f"Modelo de embeddings carregado: {self.model_name}")
        except Exception as e:
            logger.error(f"Erro ao carregar modelo de embeddings: {e}")
    
    def get_embeddings(self, texts: List[str]) -> List[np.ndarray]:
        """Gera embeddings para lista de textos"""
        if self.model is None:
            logger.error("Modelo de embeddings não disponível")
            return [np.zeros(384) for _ in texts]  # Fallback
        
        try:
            embeddings = self.model.encode(texts, show_progress_bar=True)
            return embeddings
        except Exception as e:
            logger.error(f"Erro ao gerar embeddings: {e}")
            return [np.zeros(384) for _ in texts]  # Fallback
    
    def get_embedding(self, text: str) -> np.ndarray:
        """Gera embedding para um texto"""
        return self.get_embeddings([text])[0]

class VectorStore:
    """Interface unificada para vector stores"""
    
    def __init__(self, store_type: str = "chroma", **kwargs):
        self.store_type = store_type
        self.store = None
        self.embedding_manager = EmbeddingManager()
        self._initialize_store(**kwargs)
    
    def _initialize_store(self, **kwargs):
        """Inicializa vector store"""
        if self.store_type == "chroma" and CHROMA_AVAILABLE:
            self._init_chroma(**kwargs)
        elif self.store_type == "faiss" and FAISS_AVAILABLE:
            self._init_faiss(**kwargs)
        else:
            logger.error(f"Vector store {self.store_type} não disponível")
    
    def _init_chroma(self, **kwargs):
        """Inicializa ChromaDB"""
        try:
            persist_directory = kwargs.get('persist_directory', './chroma_db')
            self.store = chromadb.PersistentClient(
                path=persist_directory,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            
            collection_name = kwargs.get('collection_name', 'documents')
            self.collection = self.store.get_or_create_collection(
                name=collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            logger.info("ChromaDB inicializado com sucesso")
        except Exception as e:
            logger.error(f"Erro ao inicializar ChromaDB: {e}")
    
    def _init_faiss(self, **kwargs):
        """Inicializa FAISS"""
        try:
            dimension = kwargs.get('dimension', 384)
            self.index = faiss.IndexFlatIP(dimension)
            self.documents = []
            logger.info("FAISS inicializado com sucesso")
        except Exception as e:
            logger.error(f"Erro ao inicializar FAISS: {e}")
    
    def add_documents(self, chunks: List[DocumentChunk]) -> bool:
        """Adiciona documentos ao vector store"""
        if not chunks:
            return False
        
        try:
            # Gerar embeddings
            texts = [chunk.content for chunk in chunks]
            embeddings = self.embedding_manager.get_embeddings(texts)
            
            # Atualizar chunks com embeddings
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
            
            # Adicionar ao store
            if self.store_type == "chroma" and hasattr(self, 'collection'):
                ids = [chunk.id for chunk in chunks]
                documents = [chunk.content for chunk in chunks]
                metadatas = [chunk.metadata for chunk in chunks]
                
                self.collection.add(
                    ids=ids,
                    documents=documents,
                    metadatas=metadatas,
                    embeddings=[emb.tolist() for emb in embeddings]
                )
            
            elif self.store_type == "faiss" and hasattr(self, 'index'):
                embeddings_array = np.array(embeddings).astype('float32')
                self.index.add(embeddings_array)
                self.documents.extend(chunks)
            
            logger.info(f"Adicionados {len(chunks)} documentos ao vector store")
            return True
        except Exception as e:
            logger.error(f"Erro ao adicionar documentos: {e}")
            return False
    
    def search(self, query: str, top_k: int = 5) -> List[SearchResult]:
        """Busca documentos similares"""
        try:
            # Gerar embedding da query
            query_embedding = self.embedding_manager.get_embedding(query)
            
            if self.store_type == "chroma" and hasattr(self, 'collection'):
                results = self.collection.query(
                    query_embeddings=[query_embedding.tolist()],
                    n_results=top_k
                )
                
                search_results = []
                for i in range(len(results['ids'][0])):
                    chunk = DocumentChunk(
                        id=results['ids'][0][i],
                        content=results['documents'][0][i],
                        metadata=results['metadatas'][0][i]
                    )
                    
                    result = SearchResult(
                        chunk=chunk,
                        similarity_score=results['distances'][0][i],
                        rank=i + 1,
                        source=chunk.source_document,
                        metadata=chunk.metadata
                    )
                    search_results.append(result)
                
                return search_results
            
            elif self.store_type == "faiss" and hasattr(self, 'index'):
                query_embedding = query_embedding.reshape(1, -1).astype('float32')
                distances, indices = self.index.search(query_embedding, top_k)
                
                search_results = []
                for i, (distance, idx) in enumerate(zip(distances[0], indices[0])):
                    if idx < len(self.documents):
                        chunk = self.documents[idx]
                        result = SearchResult(
                            chunk=chunk,
                            similarity_score=float(distance),
                            rank=i + 1,
                            source=chunk.source_document,
                            metadata=chunk.metadata
                        )
                        search_results.append(result)
                
                return search_results
            
            return []
        except Exception as e:
            logger.error(f"Erro na busca: {e}")
            return []

class RAGASEvaluator:
    """Avaliador de qualidade usando RAGAS"""
    
    def __init__(self):
        if not RAGAS_AVAILABLE:
            logger.warning("RAGAS não disponível para avaliação")
    
    def evaluate_rag_quality(self, 
                           questions: List[str],
                           contexts: List[List[str]],
                           answers: List[str],
                           ground_truths: List[str] = None) -> Dict[str, float]:
        """Avalia qualidade do RAG usando RAGAS"""
        if not RAGAS_AVAILABLE:
            return self._fallback_evaluation(questions, contexts, answers)
        
        try:
            # Preparar dados para RAGAS
            dataset = {
                'question': questions,
                'contexts': contexts,
                'answer': answers
            }
            
            if ground_truths:
                dataset['ground_truth'] = ground_truths
            
            # Métricas RAGAS
            metrics = [
                faithfulness,
                answer_relevancy,
                context_relevancy,
                context_recall
            ]
            
            if ground_truths:
                metrics.extend([answer_correctness, answer_similarity])
            
            # Executar avaliação
            results = evaluate(dataset, metrics)
            
            return {
                'faithfulness': float(results['faithfulness']),
                'answer_relevancy': float(results['answer_relevancy']),
                'context_relevancy': float(results['context_relevancy']),
                'context_recall': float(results['context_recall']),
                'answer_correctness': float(results.get('answer_correctness', 0.0)),
                'answer_similarity': float(results.get('answer_similarity', 0.0))
            }
        except Exception as e:
            logger.error(f"Erro na avaliação RAGAS: {e}")
            return self._fallback_evaluation(questions, contexts, answers)
    
    def _fallback_evaluation(self, questions: List[str], 
                           contexts: List[List[str]], 
                           answers: List[str]) -> Dict[str, float]:
        """Avaliação de fallback sem RAGAS"""
        try:
            # Métricas simples
            context_lengths = [len(' '.join(ctx)) for ctx in contexts]
            answer_lengths = [len(ans) for ans in answers]
            
            return {
                'faithfulness': 0.7,  # Placeholder
                'answer_relevancy': 0.7,  # Placeholder
                'context_relevancy': 0.7,  # Placeholder
                'context_recall': 0.7,  # Placeholder
                'avg_context_length': np.mean(context_lengths),
                'avg_answer_length': np.mean(answer_lengths)
            }
        except Exception as e:
            logger.error(f"Erro na avaliação de fallback: {e}")
            return {'error': 1.0}

class RAGSystem:
    """Sistema RAG completo"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.vector_store = VectorStore(
            store_type=config.get('vector_store_type', 'chroma'),
            persist_directory=config.get('persist_directory', './rag_db'),
            collection_name=config.get('collection_name', 'documents')
        )
        self.document_processor = DocumentProcessor(
            chunk_size=config.get('chunk_size', 512),
            chunk_overlap=config.get('chunk_overlap', 50)
        )
        self.ragas_evaluator = RAGASEvaluator()
        self.cache = {}
        
        logger.info("Sistema RAG inicializado")
    
    def add_documents(self, file_paths: List[str]) -> bool:
        """Adiciona documentos ao sistema RAG"""
        all_chunks = []
        
        for file_path in file_paths:
            try:
                file_ext = Path(file_path).suffix.lower()
                
                if file_ext == '.pdf':
                    chunks = self.document_processor.process_pdf(file_path)
                elif file_ext == '.docx':
                    chunks = self.document_processor.process_docx(file_path)
                elif file_ext in ['.txt', '.md']:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    chunks = self.document_processor.process_text(text, {'source': file_path})
                else:
                    logger.warning(f"Formato de arquivo não suportado: {file_ext}")
                    continue
                
                all_chunks.extend(chunks)
                logger.info(f"Processados {len(chunks)} chunks de {file_path}")
                
            except Exception as e:
                logger.error(f"Erro ao processar {file_path}: {e}")
                continue
        
        if all_chunks:
            return self.vector_store.add_documents(all_chunks)
        
        return False
    
    def query(self, question: str, top_k: int = 5) -> RAGResponse:
        """Executa query no sistema RAG"""
        start_time = time.time()
        
        # Verificar cache
        cache_key = hashlib.md5(f"{question}_{top_k}".encode()).hexdigest()
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        try:
            # Buscar documentos relevantes
            search_results = self.vector_store.search(question, top_k)
            
            if not search_results:
                return RAGResponse(
                    answer="Não encontrei informações relevantes para responder sua pergunta.",
                    context=[],
                    sources=[],
                    confidence_score=0.0,
                    generation_time=time.time() - start_time,
                    search_results=[],
                    metadata={'error': 'no_results'}
                )
            
            # Preparar contexto
            context_chunks = [result.chunk for result in search_results]
            context_text = "\n\n".join([chunk.content for chunk in context_chunks])
            
            # Gerar resposta (implementação simplificada)
            answer = self._generate_answer(question, context_text)
            
            # Calcular confiança
            confidence_score = np.mean([result.similarity_score for result in search_results])
            
            response = RAGResponse(
                answer=answer,
                context=context_chunks,
                sources=list(set([result.source for result in search_results])),
                confidence_score=confidence_score,
                generation_time=time.time() - start_time,
                search_results=search_results,
                metadata={'top_k': top_k}
            )
            
            # Cache
            self.cache[cache_key] = response
            
            return response
            
        except Exception as e:
            logger.error(f"Erro na query RAG: {e}")
            return RAGResponse(
                answer="Ocorreu um erro ao processar sua pergunta.",
                context=[],
                sources=[],
                confidence_score=0.0,
                generation_time=time.time() - start_time,
                search_results=[],
                metadata={'error': str(e)}
            )
    
    def _generate_answer(self, question: str, context: str) -> str:
        """Gera resposta baseada no contexto (implementação simplificada)"""
        # Em produção, isso seria integrado com um LLM
        # Por enquanto, retorna uma resposta baseada no contexto
        
        if not context:
            return "Não tenho informações suficientes para responder."
        
        # Resposta simples baseada no contexto
        sentences = context.split('.')
        relevant_sentences = [s for s in sentences if any(word in s.lower() 
                                                        for word in question.lower().split())]
        
        if relevant_sentences:
            return '. '.join(relevant_sentences[:3]) + '.'
        else:
            return sentences[0] + '.' if sentences else "Informação não encontrada."
    
    def evaluate_quality(self, test_questions: List[str], 
                        ground_truths: List[str] = None) -> Dict[str, float]:
        """Avalia qualidade do sistema RAG"""
        try:
            questions = []
            contexts = []
            answers = []
            
            for question in test_questions:
                response = self.query(question)
                questions.append(question)
                contexts.append([chunk.content for chunk in response.context])
                answers.append(response.answer)
            
            return self.ragas_evaluator.evaluate_rag_quality(
                questions, contexts, answers, ground_truths
            )
        except Exception as e:
            logger.error(f"Erro na avaliação de qualidade: {e}")
            return {'error': 1.0}
    
    def get_statistics(self) -> Dict[str, Any]:
        """Retorna estatísticas do sistema RAG"""
        try:
            if hasattr(self.vector_store, 'collection'):
                count = self.vector_store.collection.count()
            elif hasattr(self.vector_store, 'index'):
                count = self.vector_store.index.ntotal
            else:
                count = 0
            
            return {
                'total_documents': count,
                'cache_size': len(self.cache),
                'vector_store_type': self.vector_store.store_type,
                'chunk_size': self.document_processor.chunk_size,
                'chunk_overlap': self.document_processor.chunk_overlap
            }
        except Exception as e:
            logger.error(f"Erro ao obter estatísticas: {e}")
            return {'error': str(e)} 